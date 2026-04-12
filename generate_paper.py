#!/usr/bin/env python3
"""
Recreate CloudEmu Research Paper in IEEE two-column format (matching original)
and insert diagrams. Original format: two-column, black text, bold centered
section headings, italic sub-headings, italic abstract, standard academic tables.
"""

import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml

OUTPUT = "/Users/zopdev/Documents/CloudEmu/CloudEmu_Research_Paper.docx"
IMG_DIR = "/Users/zopdev/Documents/CloudEmu/_paper_imgs"
os.makedirs(IMG_DIR, exist_ok=True)

# Diagram colors
C_PRIMARY = "#1a5276"
C_AWS = "#FF9900"
C_AZURE = "#0078D4"
C_GCP = "#4285F4"
C_LAYER1 = "#2E86C1"
C_LAYER2 = "#28B463"
C_LAYER3 = "#E74C3C"
C_DARK = "#2C3E50"

def fig_to_path(fig, name):
    path = os.path.join(IMG_DIR, f"{name}.png")
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white", edgecolor="none")
    plt.close(fig)
    return path

def draw_rounded_box(ax, xy, w, h, text, color, text_color="white", fontsize=9, alpha=1.0, lw=0):
    box = FancyBboxPatch(xy, w, h, boxstyle="round,pad=0.02", facecolor=color,
                         edgecolor="white", linewidth=lw, alpha=alpha)
    ax.add_patch(box)
    cx, cy = xy[0] + w / 2, xy[1] + h / 2
    ax.text(cx, cy, text, ha="center", va="center", fontsize=fontsize,
            color=text_color, fontweight="bold", wrap=True)

def draw_arrow(ax, start, end, color="#555555", style="-|>", lw=1.5):
    ax.annotate("", xy=end, xytext=start,
                arrowprops=dict(arrowstyle=style, color=color, lw=lw))

# ── Diagram generators ────────────────────────────────────────────────────

def make_intro_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(7, 3.8))
    ax.set_xlim(0, 10); ax.set_ylim(0, 6); ax.axis("off")
    ax.text(5, 5.6, "Cloud Testing Landscape", ha="center", fontsize=13, fontweight="bold", color=C_DARK)
    for x, y, title, color, desc in [
        (0.3, 4.0, "Real Cloud\nAPIs", "#E74C3C", "Slow, $$$ Cost\nFlaky"),
        (2.7, 4.0, "Docker-Based\nEmulators", "#E67E22", "Heavy, Slow Start\nAWS Only"),
        (5.1, 4.0, "Language-Specific\nMocks", "#F39C12", "Python Only\nNo Behaviors"),
        (7.5, 4.0, "Official\nEmulators", "#D35400", "Single Service\nSingle Cloud"),
    ]:
        draw_rounded_box(ax, (x, y), 2.0, 1.1, title, color, fontsize=7.5)
        ax.text(x + 1.0, y - 0.35, desc, ha="center", fontsize=6, color="#666", style="italic")
    ax.annotate("", xy=(5, 2.9), xytext=(5, 3.4),
                arrowprops=dict(arrowstyle="-|>", color=C_PRIMARY, lw=2.5))
    ax.text(5, 3.15, "GAP", ha="center", fontsize=8, fontweight="bold", color=C_PRIMARY,
            bbox=dict(boxstyle="round,pad=0.12", facecolor="#EBF5FB", edgecolor=C_PRIMARY, lw=1))
    border = FancyBboxPatch((1.5, 1.3), 7.0, 1.4, boxstyle="round,pad=0.02",
                            facecolor="#EBF5FB", edgecolor=C_PRIMARY, linewidth=2, alpha=0.9)
    ax.add_patch(border)
    ax.text(5, 2.45, "CloudEmu", ha="center", fontsize=12, fontweight="bold", color=C_PRIMARY)
    for x, y, text, color in [
        (2.5, 1.75, "Zero\nDependencies", "#27AE60"),
        (4.2, 1.75, "Multi-Cloud\nAWS+Azure+GCP", "#2E86C1"),
        (5.9, 1.75, "Behavioral\nFidelity", "#8E44AD"),
        (7.5, 1.75, "In-Process\n~10ms", "#E67E22"),
    ]:
        draw_rounded_box(ax, (x - 0.55, y - 0.25), 1.2, 0.55, text, color, fontsize=6)
    ax.text(5, 0.8, "Fig. 1. Cloud testing landscape and the gap CloudEmu fills.",
            ha="center", fontsize=7.5, color="#333", style="italic")
    return fig_to_path(fig, "01_intro")

def make_architecture_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(7, 5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 7); ax.axis("off")
    ax.text(5, 6.7, "Three-Layer Architecture", ha="center", fontsize=13, fontweight="bold", color=C_DARK)
    border1 = FancyBboxPatch((0.5, 5.2), 9.0, 1.1, boxstyle="round,pad=0.02",
                             facecolor="#D6EAF8", edgecolor=C_LAYER1, linewidth=2)
    ax.add_patch(border1)
    ax.text(1.3, 5.95, "Layer 1: Portable API", fontsize=9, fontweight="bold", color=C_LAYER1)
    for i, cc in enumerate(["Recording", "Metrics", "Rate Limiting", "Error Injection", "Latency"]):
        draw_rounded_box(ax, (1.0 + i * 1.7, 5.3), 1.5, 0.42, cc, C_LAYER1, fontsize=6.5)
    for x in [2.5, 5.0, 7.5]:
        draw_arrow(ax, (x, 5.2), (x, 4.7), C_DARK)
    border2 = FancyBboxPatch((0.5, 3.4), 9.0, 1.2, boxstyle="round,pad=0.02",
                             facecolor="#D5F5E3", edgecolor=C_LAYER2, linewidth=2)
    ax.add_patch(border2)
    ax.text(1.3, 4.35, "Layer 2: Driver Interfaces", fontsize=9, fontweight="bold", color=C_LAYER2)
    for i, svc in enumerate(["Storage", "Compute", "Database", "Serverless", "Network",
                              "Monitor", "IAM", "DNS", "LB", "Queue"]):
        draw_rounded_box(ax, (0.7 + i * 0.88, 3.5), 0.78, 0.42, svc, C_LAYER2, fontsize=5.5)
    for x in [2.0, 5.0, 8.0]:
        draw_arrow(ax, (x, 3.4), (x, 2.9), C_DARK)
    border3 = FancyBboxPatch((0.5, 1.2), 9.0, 1.6, boxstyle="round,pad=0.02",
                             facecolor="#FADBD8", edgecolor=C_LAYER3, linewidth=2)
    ax.add_patch(border3)
    ax.text(1.3, 2.55, "Layer 3: Provider Implementations", fontsize=9, fontweight="bold", color=C_LAYER3)
    for px, py, name, color, svcs in [
        (1.0, 1.4, "AWS", C_AWS, ["S3", "EC2", "DynamoDB", "Lambda", "SQS"]),
        (3.7, 1.4, "Azure", C_AZURE, ["Blob", "VMs", "CosmosDB", "Func", "SvcBus"]),
        (6.4, 1.4, "GCP", C_GCP, ["GCS", "GCE", "Firestore", "CloudFn", "PubSub"]),
    ]:
        draw_rounded_box(ax, (px, py), 2.5, 1.0, "", color, alpha=0.85)
        ax.text(px + 1.25, py + 0.7, name, ha="center", fontsize=9, fontweight="bold", color="white")
        for j, s in enumerate(svcs):
            ax.text(px + 0.1 + j * 0.48 + 0.2, py + 0.2, s, ha="center", fontsize=5, color="white", alpha=0.9)
    draw_rounded_box(ax, (3.0, 0.45), 4.0, 0.5, "memstore.Store[V]  (Generic Thread-Safe Map)", "#34495E", fontsize=7.5)
    ax.text(5, 0.05, "Fig. 2. Three-layer architecture with provider implementations.",
            ha="center", fontsize=7.5, color="#333", style="italic")
    return fig_to_path(fig, "02_architecture")

def make_cross_service_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(7, 4))
    ax.set_xlim(0, 10); ax.set_ylim(0, 6); ax.axis("off")
    ax.text(5, 5.7, "Cross-Service Wiring", ha="center", fontsize=13, fontweight="bold", color=C_DARK)
    for cx, name, color, svcs in [
        (0.5, "AWS", C_AWS, ["EC2", "CloudWatch", "SQS", "Lambda"]),
        (3.5, "Azure", C_AZURE, ["VMs", "Monitor", "ServiceBus", "Functions"]),
        (6.5, "GCP", C_GCP, ["GCE", "Monitoring", "Pub/Sub", "CloudFn"]),
    ]:
        draw_rounded_box(ax, (cx, 4.9), 3.0, 0.5, name, color, fontsize=9)
        draw_rounded_box(ax, (cx + 0.1, 4.0), 1.3, 0.5, svcs[0], "#5DADE2", fontsize=7.5)
        draw_rounded_box(ax, (cx + 1.6, 4.0), 1.3, 0.5, svcs[1], "#48C9B0", fontsize=7.5)
        draw_arrow(ax, (cx + 1.4, 4.25), (cx + 1.6, 4.25), "#E74C3C", lw=2)
        ax.text(cx + 1.5, 3.55, "SetMonitoring()", ha="center", fontsize=5.5, color="#E74C3C", fontweight="bold")
        draw_rounded_box(ax, (cx + 0.1, 2.5), 1.3, 0.5, svcs[2], "#AF7AC5", fontsize=7.5)
        draw_rounded_box(ax, (cx + 1.6, 2.5), 1.3, 0.5, svcs[3], "#F5B041", fontsize=7.5)
        draw_arrow(ax, (cx + 1.4, 2.75), (cx + 1.6, 2.75), "#E74C3C", lw=2)
        ax.text(cx + 1.5, 2.05, "SetTrigger()", ha="center", fontsize=5.5, color="#E74C3C", fontweight="bold")
    ax.text(5, 1.3, "Fig. 3. Cross-service wiring across providers.",
            ha="center", fontsize=7.5, color="#333", style="italic")
    return fig_to_path(fig, "03_cross_service")

def make_implementation_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(7, 4.5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 6.5); ax.axis("off")
    ax.text(5, 6.2, "Implementation Components", ha="center", fontsize=13, fontweight="bold", color=C_DARK)
    for x, y, title, color, desc in [
        (0.5, 4.5, "memstore\nStore[V]", "#2E86C1", "Generic thread-safe\nmap[string]V"),
        (2.7, 4.5, "statemachine\nFSM", "#27AE60", "VM lifecycle\ntransitions"),
        (4.9, 4.5, "pagination\nPaginate[T]", "#8E44AD", "Base64 page\ntokens"),
        (7.1, 4.5, "idgen\nIDs", "#E67E22", "ARN / Azure ID\n/ GCP SelfLink"),
    ]:
        draw_rounded_box(ax, (x, y), 2.0, 0.85, title, color, fontsize=7.5)
        ax.text(x + 1.0, y - 0.3, desc, ha="center", fontsize=6, color="#666")
    for x, y, title, color, desc in [
        (0.5, 2.8, "config\nOptions", "#1ABC9C", "WithClock, WithRegion\nWithAccountID"),
        (2.7, 2.8, "config\nFakeClock", "#16A085", "Deterministic time\nfor testing"),
        (4.9, 2.8, "errors\nCodes", "#C0392B", "NotFound, AlreadyExists\nThrottled, etc."),
        (7.1, 2.8, "cost\nTracker", "#D4AC0D", "Per-operation\ncost simulation"),
    ]:
        draw_rounded_box(ax, (x, y), 2.0, 0.85, title, color, fontsize=7.5)
        ax.text(x + 1.0, y - 0.3, desc, ha="center", fontsize=6, color="#666")
    for x, y, title, color in [
        (1.5, 1.2, "Recorder", "#5B2C6F"), (3.3, 1.2, "Metrics\nCollector", "#1A5276"),
        (5.1, 1.2, "Rate\nLimiter", "#7B241C"), (6.9, 1.2, "Error\nInjector", "#784212"),
    ]:
        draw_rounded_box(ax, (x, y), 1.5, 0.65, title, color, fontsize=7)
    ax.text(5, 0.55, "Fig. 4. Internal implementation components.",
            ha="center", fontsize=7.5, color="#333", style="italic")
    return fig_to_path(fig, "04_implementation")

def make_vm_lifecycle_flowchart():
    fig, ax = plt.subplots(1, 1, figsize=(7, 3.5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 5); ax.axis("off")
    ax.text(5, 4.7, "VM Lifecycle State Machine", ha="center", fontsize=13, fontweight="bold", color=C_DARK)
    states = {
        "pending": (1.2, 3.0, "#F39C12"), "running": (3.5, 3.0, "#27AE60"),
        "stopping": (5.8, 3.8, "#E67E22"), "stopped": (8.0, 3.0, "#95A5A6"),
        "shutting-down": (5.8, 2.2, "#E74C3C"), "terminated": (8.0, 1.5, "#C0392B"),
    }
    for name, (x, y, color) in states.items():
        draw_rounded_box(ax, (x - 0.65, y - 0.22), 1.3, 0.45, name, color, fontsize=7)
    for src, dst, label in [
        ("pending", "running", "auto"), ("running", "stopping", "Stop"),
        ("stopping", "stopped", "auto"), ("stopped", "pending", "Start"),
        ("running", "shutting-down", "Terminate"), ("shutting-down", "terminated", "auto"),
    ]:
        sx, sy, _ = states[src]; dx, dy, _ = states[dst]
        if sy == dy:
            draw_arrow(ax, (sx + 0.6, sy), (dx - 0.7, dy), C_DARK, lw=1.5)
            ax.text((sx + dx) / 2, sy + 0.32, label, ha="center", fontsize=6, color=C_PRIMARY, fontweight="bold")
        else:
            draw_arrow(ax, (sx + 0.5, sy + (0.08 if dy > sy else -0.22)),
                       (dx - 0.5, dy + (-0.08 if dy > sy else 0.22)), C_DARK, lw=1.5)
            ax.text((sx + dx) / 2 + 0.3, (sy + dy) / 2, label, fontsize=6, color=C_PRIMARY, fontweight="bold")
    rx, ry, _ = states["running"]
    ax.annotate("", xy=(rx + 0.2, ry + 0.27), xytext=(rx - 0.2, ry + 0.27),
                arrowprops=dict(arrowstyle="-|>", color="#8E44AD", lw=1.5, connectionstyle="arc3,rad=-0.8"))
    ax.text(rx, ry + 0.7, "Reboot", ha="center", fontsize=6, color="#8E44AD", fontweight="bold")
    ax.text(5, 0.85, "Fig. 5. VM lifecycle state machine.",
            ha="center", fontsize=7.5, color="#333", style="italic")
    return fig_to_path(fig, "05_flowchart_vm")

def make_alarm_flowchart():
    fig, ax = plt.subplots(1, 1, figsize=(7, 4.2))
    ax.set_xlim(0, 10); ax.set_ylim(0, 6.5); ax.axis("off")
    ax.text(5, 6.2, "Auto-Metric and Alarm Evaluation Flow", ha="center",
            fontsize=13, fontweight="bold", color=C_DARK)
    steps = [
        (1.5, 5.0, "RunInstances()\ncalled", "#2E86C1"),
        (1.5, 3.8, "emitInstance\nMetrics()", "#27AE60"),
        (4.5, 3.8, "PutMetricData()\n5 metrics x 5 pts", "#8E44AD"),
        (7.5, 3.8, "evaluateAlarms()\ntriggered", "#E74C3C"),
        (7.5, 2.5, "Collect datapoints\nin eval window", "#F39C12"),
        (7.5, 1.3, "Compare statistic\nvs threshold", "#E67E22"),
        (4.5, 1.3, "Update alarm\nstate (OK/ALARM)", "#C0392B"),
        (1.5, 1.3, "Alarm state\nchanged!", "#27AE60"),
    ]
    for x, y, text, color in steps:
        draw_rounded_box(ax, (x - 0.75, y - 0.3), 1.7, 0.65, text, color, fontsize=7)
    for i, j in [(0,1),(1,2),(2,3),(3,4),(4,5),(5,6),(6,7)]:
        sx, sy = steps[i][0], steps[i][1]; dx, dy = steps[j][0], steps[j][1]
        if sx == dx:
            draw_arrow(ax, (sx, sy - 0.3), (dx, dy + 0.35), C_DARK, lw=1.5)
        elif dx > sx:
            draw_arrow(ax, (sx + 0.85, sy), (dx - 0.85, dy), C_DARK, lw=1.5)
        else:
            draw_arrow(ax, (sx - 0.85, sy), (dx + 0.85, dy), C_DARK, lw=1.5)
    ax.text(5, 0.55, "Fig. 6. Auto-metric generation and alarm evaluation flow.",
            ha="center", fontsize=7.5, color="#333", style="italic")
    return fig_to_path(fig, "06_alarm_flow")

def make_comparison_bar_chart():
    fig, axes = plt.subplots(1, 3, figsize=(8, 3))
    tools = ["CloudEmu", "LocalStack", "Moto", "Azurite", "Firebase\nEmulator"]
    colors = [C_PRIMARY, "#E74C3C", "#F39C12", C_AZURE, "#FBBC04"]
    startup = [0, 3000, 500, 2000, 4000]
    axes[0].barh(tools, startup, color=colors, edgecolor="white", height=0.6)
    axes[0].set_xlabel("Startup Time (ms)", fontsize=7)
    axes[0].set_title("Startup Time", fontsize=9, fontweight="bold")
    axes[0].tick_params(labelsize=6)
    for i, v in enumerate(startup):
        axes[0].text(v + 80, i, f"{v}ms", va="center", fontsize=6)
    providers = [3, 1, 1, 1, 1]
    axes[1].barh(tools, providers, color=colors, edgecolor="white", height=0.6)
    axes[1].set_xlabel("Cloud Providers", fontsize=7)
    axes[1].set_title("Provider Coverage", fontsize=9, fontweight="bold")
    axes[1].tick_params(labelsize=6)
    for i, v in enumerate(providers):
        axes[1].text(v + 0.05, i, str(v), va="center", fontsize=6)
    deps = [0, 2, 1, 1, 2]
    dep_labels = ["None", "Docker+Py", "Python", "Node.js", "Node+Java"]
    axes[2].barh(tools, deps, color=colors, edgecolor="white", height=0.6)
    axes[2].set_xlabel("External Dependencies", fontsize=7)
    axes[2].set_title("Dependencies", fontsize=9, fontweight="bold")
    axes[2].tick_params(labelsize=6)
    for i, v in enumerate(deps):
        axes[2].text(v + 0.05, i, dep_labels[i], va="center", fontsize=5.5)
    plt.tight_layout()
    fig.text(0.5, -0.01, "Fig. 7. Comparison of CloudEmu with existing cloud testing tools.",
             ha="center", fontsize=7.5, color="#333", style="italic")
    return fig_to_path(fig, "07_comparison_bar")

def make_comparison_radar():
    categories = ["Zero\nDependency", "Multi-Cloud", "Behavioral\nFidelity",
                   "Speed", "Language\nAgnostic", "Service\nBreadth"]
    N = len(categories)
    angles = [n / float(N) * 2 * np.pi for n in range(N)]
    angles += angles[:1]
    tools_data = {
        "CloudEmu": [10,10,9,10,3,8], "LocalStack": [2,2,8,4,7,10],
        "Moto": [5,2,6,7,2,8], "Azurite": [4,2,7,5,6,2],
    }
    tool_colors = {"CloudEmu": C_PRIMARY, "LocalStack": "#E74C3C", "Moto": "#F39C12", "Azurite": C_AZURE}
    fig, ax = plt.subplots(1, 1, figsize=(5, 4.2), subplot_kw=dict(polar=True))
    ax.set_theta_offset(np.pi / 2); ax.set_theta_direction(-1); ax.set_rlabel_position(0)
    plt.xticks(angles[:-1], categories, fontsize=7)
    plt.yticks([2,4,6,8,10], ["2","4","6","8","10"], fontsize=6, color="grey")
    plt.ylim(0, 11)
    for tool, values in tools_data.items():
        vals = values + values[:1]
        ax.plot(angles, vals, linewidth=1.8, label=tool, color=tool_colors[tool])
        ax.fill(angles, vals, alpha=0.08, color=tool_colors[tool])
    ax.legend(loc="upper right", bbox_to_anchor=(1.3, 1.1), fontsize=7)
    ax.set_title("Multi-Dimensional Comparison", fontsize=11, fontweight="bold", color=C_DARK, pad=15)
    fig.text(0.5, 0.01, "Fig. 8. Radar comparison across key dimensions.",
             ha="center", fontsize=7.5, color="#333", style="italic")
    return fig_to_path(fig, "08_comparison_radar")

def make_coverage_heatmap():
    fig, ax = plt.subplots(1, 1, figsize=(6.5, 3.5))
    services = ["Storage", "Compute", "Database", "Serverless", "Network",
                "Monitor", "IAM", "DNS", "LB", "MsgQ"]
    tools_list = ["CloudEmu\nAWS", "CloudEmu\nAzure", "CloudEmu\nGCP",
                  "LocalStack", "Moto", "Azurite", "Firebase\nEmulator"]
    data = np.array([
        [1,1,1,1,1,1,1,1,1,1],[1,1,1,1,1,1,1,1,1,1],[1,1,1,1,1,1,1,1,1,1],
        [1,1,1,1,1,1,1,1,1,1],[1,1,1,1,0,0,1,1,0,1],[1,0,0,0,0,0,0,0,0,0],
        [0,0,1,1,0,0,1,0,0,0],
    ])
    cmap = plt.cm.colors.ListedColormap(["#FADBD8", "#82E0AA"])
    ax.imshow(data, cmap=cmap, aspect="auto")
    ax.set_xticks(range(len(services))); ax.set_xticklabels(services, fontsize=6, rotation=45, ha="right")
    ax.set_yticks(range(len(tools_list))); ax.set_yticklabels(tools_list, fontsize=7)
    for i in range(len(tools_list)):
        for j in range(len(services)):
            sym = "\u2713" if data[i,j]==1 else "\u2717"
            clr = "#1a5276" if data[i,j]==1 else "#C0392B"
            ax.text(j, i, sym, ha="center", va="center", fontsize=9, color=clr, fontweight="bold")
    ax.set_title("Service Coverage Comparison", fontsize=11, fontweight="bold", color=C_DARK)
    for i in range(len(tools_list)+1): ax.axhline(i-0.5, color="white", linewidth=2)
    for j in range(len(services)+1): ax.axvline(j-0.5, color="white", linewidth=2)
    ax.axhline(2.5, color=C_PRIMARY, linewidth=2.5)
    plt.tight_layout()
    fig.text(0.5, -0.01, "Fig. 9. Service coverage heatmap.",
             ha="center", fontsize=7.5, color="#333", style="italic")
    return fig_to_path(fig, "09_coverage_heatmap")

# Generate all diagrams
print("Generating diagrams...")
img_intro = make_intro_diagram()
img_arch = make_architecture_diagram()
img_cross = make_cross_service_diagram()
img_impl = make_implementation_diagram()
img_vm = make_vm_lifecycle_flowchart()
img_alarm = make_alarm_flowchart()
img_bar = make_comparison_bar_chart()
img_radar = make_comparison_radar()
img_coverage = make_coverage_heatmap()
print("All diagrams generated.")

# ══════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT — IEEE two-column format matching original exactly
# ══════════════════════════════════════════════════════════════════════════

doc = Document()

# Set page to Letter size with IEEE margins
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(1.57)
    section.right_margin = Cm(1.57)

# Enable two-column layout
sectPr = doc.sections[0]._sectPr
cols = parse_xml(f'<w:cols {nsdecls("w")} w:num="2" w:space="720"/>')
sectPr.append(cols)

# ── Style helpers matching original IEEE format ──
def p_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(24)
    r.font.name = "Times New Roman"
    return p

def p_author(doc, text, bold=True, size=Pt(10)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = size
    r.font.name = "Times New Roman"
    return p

def p_section(doc, text):
    """Centered bold uppercase section heading — e.g. 'III. ARCHITECTURE AND DESIGN'"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    r.font.name = "Times New Roman"
    return p

def p_subsection(doc, text):
    """Left-aligned italic sub-heading — e.g. 'A. Design Goals'"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    r.font.name = "Times New Roman"
    return p

def p_body(doc, text, indent_first=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(0.5)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = "Times New Roman"
    return p

def p_abstract_body(doc, text):
    """Abstract text in italic per IEEE style."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    r.font.name = "Times New Roman"
    return p

def p_keywords(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Keywords\u2014")
    r.bold = True
    r.italic = True
    r.font.size = Pt(10)
    r.font.name = "Times New Roman"
    r2 = p.add_run(text)
    r2.italic = True
    r2.font.size = Pt(10)
    r2.font.name = "Times New Roman"
    return p

def p_bullet_bold_start(doc, bold_text, rest_text):
    """Bullet point with bold start — matching original format."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.clear()
    r1 = p.add_run(bold_text)
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.name = "Times New Roman"
    r2 = p.add_run(rest_text)
    r2.font.size = Pt(10)
    r2.font.name = "Times New Roman"
    return p

def p_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.clear()
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = "Times New Roman"
    return p

def table_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(8)
    r.font.name = "Times New Roman"
    r.bold = True
    return p

def add_ieee_table(doc, headers, rows, col_widths=None):
    """Standard academic table — bold headers, thin borders, fits in one column.
    Uses XML-level width constraints so Word respects column boundaries."""
    TABLE_W = 4600  # total table width in twips (~3.19 inches, fits one column)
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = False

    # ── Force table width at XML level ──
    tbl = table._tbl
    tblPr = tbl.tblPr
    # Set fixed table width
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{TABLE_W}" w:type="dxa"/>')
    # Remove any existing tblW
    for existing in tblPr.findall(qn('w:tblW')):
        tblPr.remove(existing)
    tblPr.append(tblW)
    # Set layout to fixed (prevents auto-expand)
    tblLayout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
    for existing in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(existing)
    tblPr.append(tblLayout)

    # ── Set grid column widths at XML level ──
    if col_widths is None:
        cw_twips = [TABLE_W // n_cols] * n_cols
    else:
        cw_twips = [int(w) for w in col_widths]
    # Remove existing tblGrid and create new
    for existing in tbl.findall(qn('w:tblGrid')):
        tbl.remove(existing)
    grid_xml = f'<w:tblGrid {nsdecls("w")}>'
    for cw in cw_twips:
        grid_xml += f'<w:gridCol w:w="{cw}"/>'
    grid_xml += '</w:tblGrid>'
    tbl.insert(tbl.index(tblPr) + 1, parse_xml(grid_xml))

    # ── Set each cell width explicitly ──
    for row_obj in table.rows:
        for j, cell in enumerate(row_obj.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{cw_twips[j]}" w:type="dxa"/>')
            for existing in tcPr.findall(qn('w:tcW')):
                tcPr.remove(existing)
            tcPr.append(tcW)

    # ── Header row ──
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(h)
        r.font.size = Pt(7)
        r.font.name = "Times New Roman"
        r.bold = True
    # ── Data rows ──
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            r = p.add_run(str(val))
            r.font.size = Pt(7)
            r.font.name = "Times New Roman"
    return table

def add_img(doc, path, width=Inches(3.2)):
    """Insert image centered, sized for column width."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run()
    r.add_picture(path, width=width)
    return p

def p_ref(doc, text, keep_with_next=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = keep_with_next
    r = p.add_run(text)
    r.font.size = Pt(7)
    r.font.name = "Times New Roman"
    return p


# ══════════════════════════════════════════════════════════════════════════
# DOCUMENT CONTENT — exact original text, IEEE two-column format
# ══════════════════════════════════════════════════════════════════════════

# TITLE
p_title(doc, "CloudEmu: A Zero-Dependency, In-Memory Multi-Cloud Service Emulation Library for Go")

# AUTHORS
p_author(doc, "Nitin Kumar, Gursewak Singh, Devansh Arora")
p_author(doc, "Department of Computer Science, Chitkara University,", bold=False, size=Pt(10))
p_author(doc, "Rajpura, Punjab, India", bold=False, size=Pt(10))
p_author(doc, "nitin2001.be22@chitkara.edu.in,", bold=False, size=Pt(9))
p_author(doc, "devansh1490.be22@chitkara.edu.in,", bold=False, size=Pt(9))
p_author(doc, "gursewak1601.be22@chitkara.edu.in", bold=False, size=Pt(9))

# ABSTRACT
p_section(doc, "ABSTRACT")
p_abstract_body(doc,
    "Anyone who has tried to write integration tests against live cloud endpoints knows the drill: requests crawl over the network, "
    "the monthly bill creeps up, and every other CI run fails for reasons nobody can reproduce locally. Existing workarounds each "
    "come with a catch. LocalStack pulls in Docker. Moto only speaks Python and only speaks AWS. The vendor-provided emulators "
    "cover one service apiece from one provider. We grew frustrated enough to write our own solution. CloudEmu is a Go library "
    "that recreates 10 categories of cloud services for AWS, Azure, and GCP entirely in-process, with zero external dependencies. "
    "It is organized in three layers\u2014portable APIs, driver contracts, provider backends\u2014but honestly, the architecture is the "
    "boring part. What took us the most time was getting the behaviors right. The library pushes monitoring metrics when a virtual "
    "machine starts or stops, checks alarm thresholds against incoming data, parses IAM policy JSON and evaluates wildcard matches, "
    "silently drops duplicate FIFO messages within a five-minute window, compares database filter values numerically instead of "
    "alphabetically, routes poison messages to dead-letter queues, tallies rough per-call cost estimates, and fires serverless "
    "functions whenever a new message appears on a queue. Thirty-one tests cover all of it; they complete in under half a second "
    "and we have yet to see a single non-deterministic failure.")
p_keywords(doc, "cloud testing, service emulation, test doubles, multi-cloud, in-memory fakes, Go")

# I. INTRODUCTION
p_section(doc, "I. INTRODUCTION")
p_body(doc,
    "Take any production codebase and start listing what it talks to: an object store, a couple of virtual machines, at least one "
    "database, a message queue or two, DNS records, monitoring dashboards, a load balancer in front, IAM roles gluing permissions "
    "together, maybe a Lambda function triggered by an event. Now imagine the company runs two or even three cloud providers for "
    "redundancy or compliance. How do you test all those moving parts without actually hitting the cloud? That was the question "
    "we kept running into.")
p_body(doc,
    "One option is to test against live cloud endpoints, but that introduces network latency, accumulates real charges, and produces "
    "flaky failures whenever shared state collides or an API throttles you. LocalStack [1] eliminates the bill, sure, but it pulls "
    "in Docker, takes a few seconds to boot, and only covers AWS. Moto [2] intercepts boto3 calls neatly \u2014 unless your team writes "
    "Go, in which case it is irrelevant. Official emulators like Azurite, DynamoDB Local, and the Firebase suite each tackle one "
    "service from one provider. Stitching three of them together for a single integration test is nobody\u2019s idea of fun.",
    indent_first=True)
p_body(doc,
    "And none of them deal with the multi-cloud angle at all. A growing number of engineering teams spread workloads across two or "
    "three providers \u2014 sometimes for redundancy, sometimes because procurement or compliance forces their hand. If you want to test "
    "that kind of code locally, you are on your own: cobble together one emulator per provider and hope they play nice. Often they don\u2019t.",
    indent_first=True)

# >>> DIAGRAM: Intro
add_img(doc, img_intro)

p_body(doc, "We spent a fair amount of time looking at what was already out there. Eventually we boiled down what was missing into four requirements, and as far as we could tell, nothing on the market satisfies all four at once:")
p_bullet_bold_start(doc, "No external dependencies at all. ", "No Docker daemon running, no network calls, no JVM, not even a separate process.")
p_bullet_bold_start(doc, "One library for three clouds. ", "AWS, Azure, and GCP services accessible through a single import.")
p_bullet_bold_start(doc, "Behaviors that actually match production. ", "Alarms evaluate themselves, IAM policies get parsed, FIFO dedup is enforced.")
p_bullet_bold_start(doc, "Starts instantly, runs in-process. ", "Good enough for unit tests that need to finish in milliseconds, not seconds.")

p_body(doc,
    "CloudEmu grew out of that frustration. It is an open-source Go library that recreates 10 services for each of three "
    "providers \u2014 30 service implementations in total \u2014 using a three-layer design where portable APIs sit above driver "
    "contracts, which in turn sit above provider-specific in-memory backends.")

p_body(doc, "To summarize what this paper actually contributes:", indent_first=True)
p_bullet_bold_start(doc, "Architecture. ", "Three layers that keep portable APIs, driver contracts, and provider code from leaking into each other. Adding a fourth provider later should not require touching the top two layers.")
p_bullet_bold_start(doc, "Behavioral depth. ", "Eight emulations that go well past record storage: metrics that appear when VMs boot, alarms that trip on threshold crossings, IAM policies that actually get enforced, and five more described in Section V.")
p_bullet_bold_start(doc, "Validation. ", "Thirty-one tests across every service, including cross-service scenarios like VM-launch-triggers-alarm. The full suite runs in under 500 ms.")

# II. BACKGROUND
p_section(doc, "II. BACKGROUND")

p_subsection(doc, "A. The Cloud Testing Problem")
p_body(doc,
    "Most cloud-native code never touches a server directly; it calls managed services through vendor SDKs. A systematic review "
    "from 2019 [3] looked at 147 primary studies spanning six research areas and pointed out that lightweight, portable test tooling "
    "remains scarce. On the practitioner side, Kumar [4] measured that teams running continuous performance checks in CI/CD caught "
    "76 percent of problems before code reached production, versus only 31 percent for teams that tested manually and infrequently. "
    "Numbers like those make a convincing case for keeping cloud tests fast and frictionless.")

p_subsection(doc, "B. Test Double Taxonomy")
p_body(doc,
    "Meszaros [5] carved test doubles into five kinds: Dummy, Stub, Spy, Mock, and Fake. Fowler [6] later observed that Fakes are "
    "the most labor-intensive to write but also the most faithful stand-ins. CloudEmu belongs firmly in that Fake category. Its "
    "emulations remember state between calls, reject illegal transitions, parse and enforce policy documents, and trigger side effects "
    "across services. That is a different league from Stubs that return canned values or Mocks that just record which methods were called.")

p_subsection(doc, "C. The Multi-Cloud Gap")
p_body(doc,
    "Look at what is available and you see tunnel vision everywhere. LocalStack [1] and Moto [2] only speak AWS. Azurite [7] handles "
    "Azure Storage and nothing more. The Firebase Emulator Suite [8] is GCP-exclusive. Libraries like Go CDK [9], Libcloud [10], and "
    "CloudBridge [11] do abstract across providers, but they still delegate to real SDKs under the hood, which means you need live "
    "credentials and network access even in a test.")

# III. RELATED WORK
p_section(doc, "III. RELATED WORK")

p_subsection(doc, "A. Cloud Infrastructure Simulators")
p_body(doc,
    "CloudSim [12], CloudSim Plus [13], and CloudSim 7G [14] target a fundamentally different audience. Their focus is data-center "
    "simulation \u2014 modeling hosts, VM placement, network links \u2014 for research on scheduling algorithms and resource allocation. "
    "They do not expose the kind of service-level API an application developer would call. You cannot invoke CreateBucket() inside "
    "CloudSim; that concept does not exist there.")

p_subsection(doc, "B. Cloud Service Emulators")
p_body(doc,
    "LocalStack [1] comes closest to what we were after \u2014 over 70 AWS services with fairly high fidelity \u2014 but it drags in Docker "
    "and only talks AWS. Moto [2] hooks into Python\u2019s boto3 at the transport layer, which is slick but useless if your codebase is "
    "Go. The official single-service emulators (Azurite [7], DynamoDB Local, Firebase Suite [8]) each cover one narrow slice of one cloud.")

p_subsection(doc, "C. Cloud Abstraction Libraries")
p_body(doc,
    "Go CDK [9] gives you portable Go APIs, which sounds promising until you realize it still delegates every call to the real "
    "vendor SDK \u2014 so live credentials and network access remain non-negotiable. CloudBridge [11], jclouds [15], and Libcloud [10] "
    "follow the identical pattern, just in Python or Java.")

p_subsection(doc, "D. Service Virtualization")
p_body(doc,
    "WireMock [16], Hoverfly, and Mountebank operate one level lower \u2014 they capture HTTP traffic and replay it. That works well "
    "for regression suites, but a recorded response cannot enforce IAM rules, evaluate alarm thresholds, or reject a duplicate "
    "FIFO message. The logic simply is not there.")

p_subsection(doc, "E. Testing Methodology")
p_body(doc,
    "The test-double vocabulary introduced by Fowler [6] and Meszaros [5] is still the industry standard. Wang et al. [17] examined "
    "Apache projects and found mocking in 66 percent of them, yet almost every case turned out to be Stubs or Mocks rather than Fakes. "
    "The reason is obvious: Fakes are expensive to build. CloudEmu absorbs that cost once so that individual teams do not have to "
    "reinvent the same in-memory implementations project after project.")

# >>> DIAGRAM: Coverage heatmap
add_img(doc, img_coverage)

# III. ARCHITECTURE AND DESIGN (renumbered to IV due to Related Work move)
p_section(doc, "IV. ARCHITECTURE AND DESIGN")

p_subsection(doc, "A. Design Goals")
p_bullet_bold_start(doc, "Zero dependencies. ", "If it does not ship with Go\u2019s standard library, we do not import it. No Docker, no JVM, no pip.")
p_bullet_bold_start(doc, "Behavioral fidelity. ", "Returning canned data was not enough. We wanted state machines that reject bad transitions, policy documents that actually get parsed, alarms that actually fire when thresholds cross.")
p_bullet_bold_start(doc, "Thread safety. ", "Go programs spin up goroutines freely, so every operation protects shared state with sync.RWMutex. We ran race-detector builds throughout development.")
p_bullet_bold_start(doc, "Deterministic testing. ", "A pluggable FakeClock replaces time.Now(). Tests produce the same output whether they run on a fast CI box or a slow laptop at midnight.")
p_bullet_bold_start(doc, "Uniform provider coverage. ", "AWS, Azure, and GCP each implement the same ten categories. We refused to ship GCP with half the features of AWS.")

p_subsection(doc, "B. Three-Layer Architecture")
p_body(doc,
    "The design borrows shamelessly from Go\u2019s database/sql package, which separates the interface consumers use from the driver implementations behind it. We liked that split but found we also needed a third layer on top for recording, rate limiting, and the other cross-cutting concerns (Table I).")

# >>> DIAGRAM: Architecture
add_img(doc, img_arch)

table_caption(doc, "TABLE I. THREE-LAYER ARCHITECTURE")
add_ieee_table(doc,
    ["Layer", "Location", "Purpose"],
    [
        ["Portable API", "storage/, compute/, ...", "Recording, metrics, rate limiting"],
        ["Driver Interfaces", "*/driver/driver.go", "Go interfaces per service"],
        ["Provider Impls", "providers/{aws,azure,gcp}/", "In-memory memstore backends"],
    ],
    col_widths=[1300, 1700, 1600])

p_subsection(doc, "C. Service Coverage")
p_body(doc,
    "Table II lists the ten categories every provider implements. In practice, switching a test from AWS to GCP is surprisingly easy: change NewAWS() to NewGCP(), fix a few assertion strings that mention provider-specific IDs, and you are done. At least, that has been our experience so far.")

table_caption(doc, "TABLE II. SERVICE COVERAGE ACROSS THREE CLOUD PROVIDERS")
add_ieee_table(doc,
    ["Category", "AWS", "Azure", "GCP"],
    [
        ["Storage", "S3", "Blob Storage", "GCS"],
        ["Compute", "EC2", "VMs", "GCE"],
        ["Database", "DynamoDB", "CosmosDB", "Firestore"],
        ["Serverless", "Lambda", "Functions", "CloudFn"],
        ["Network", "VPC", "VNet", "VPC"],
        ["Monitoring", "CloudWatch", "Monitor", "Monitoring"],
        ["IAM", "IAM", "Azure IAM", "Cloud IAM"],
        ["DNS", "Route 53", "Azure DNS", "Cloud DNS"],
        ["Load Bal.", "ELB", "Azure LB", "Cloud LB"],
        ["Msg Queue", "SQS", "Service Bus", "Pub/Sub"],
    ],
    col_widths=[1100, 1100, 1200, 1200])

p_subsection(doc, "D. Cross-Service Wiring")
p_body(doc,
    "In a real AWS account, services are not isolated islands. Start an EC2 instance and CloudWatch begins collecting CPU and network "
    "data without anyone asking it to. We wanted the same thing in our fakes. So when the provider factory initializes, it calls "
    "SetMonitoring() on the compute mock, handing it a reference to the monitoring mock. After that, spinning up a VM automatically "
    "generates metrics, which in turn can trip an alarm\u2014all without the test author writing a single line of metric-injection code. "
    "Getting that wiring right took more iteration than we expected.")

# >>> DIAGRAM: Cross-service wiring
add_img(doc, img_cross)

p_subsection(doc, "E. Deterministic Time")
p_body(doc,
    "Anywhere the code needs the current time, it asks a Clock interface instead of calling time.Now() directly. Tests inject a "
    "FakeClock and advance it by arbitrary amounts. This sounds like a minor detail, but it is absolutely essential. Without it, "
    "testing a five-minute deduplication window would mean actually waiting five minutes, which is obviously unacceptable in a "
    "unit test.")

# IV. REALISTIC CLOUD BEHAVIORS
p_section(doc, "V. REALISTIC CLOUD BEHAVIORS")
p_body(doc,
    "We put substantial effort into eight behaviors that distinguish a useful fake from a glorified key-value store. Every one of "
    "them works the same way regardless of whether you instantiate the AWS, Azure, or GCP provider.")

p_subsection(doc, "A. Auto-Metric Generation on Compute Lifecycle")
p_body(doc,
    "Call RunInstances and five metrics appear in the monitoring mock automatically. Stop the VM and every metric falls to zero. "
    "Start it again and they jump back up. Reboot, terminate \u2014 each lifecycle action pushes its own datapoints. A test can set up "
    "an alarm, launch a VM, and verify the alarm tripped, all without touching the monitoring API directly. One quirk we had to "
    "handle: GCP reports CPU utilization as a fraction (0.25) whereas AWS and Azure report it as a percentage (25.0).")

table_caption(doc, "TABLE III. METRIC EMISSION PER LIFECYCLE OP")
add_ieee_table(doc,
    ["Operation", "CPU", "Net", "Disk", "Pts"],
    [
        ["Run", "25.0", "1024/512", "100/50", "5"],
        ["Start", "25.0", "1024/512", "100/50", "1"],
        ["Stop", "0.0", "0.0/0.0", "0.0/0.0", "1"],
        ["Reboot", "25.0", "1024/512", "100/50", "1"],
        ["Terminate", "0.0", "0.0/0.0", "0.0/0.0", "1"],
    ],
    col_widths=[1100, 650, 850, 850, 550])

p_subsection(doc, "B. Alarm Auto-Evaluation")
p_body(doc,
    "Whenever metric data arrives, the monitoring mock scans its alarm list for any alarm that watches the same namespace and metric. "
    "Matching alarms pull datapoints that fall within their evaluation window \u2014 Period multiplied by EvaluationPeriods \u2014 compute "
    "the requested statistic (average, sum, min, max, or sample count), and compare the result against the threshold. If that "
    "comparison crosses the boundary, the alarm flips state. We hit a deadlock during development here: the alarm evaluation "
    "tries to read metric data, but the write lock from PutMetricData was still held. The fix was to release the write lock "
    "before calling the evaluator.")

# >>> DIAGRAM: Alarm flow
add_img(doc, img_alarm)

p_subsection(doc, "C. IAM Policy Evaluation")
p_body(doc,
    "CheckPermission(principal, action, resource) collects every policy document attached to the given principal, including both "
    "user-attached and role-attached policies. Each document\u2019s JSON gets parsed, and the requested action and resource are compared "
    "against every statement via wildcardMatch(), a helper that handles * and ? patterns. The evaluation logic mirrors real IAM: "
    "an explicit Deny anywhere trumps everything. Absent a Deny, at least one Allow must match. If nothing matches at all, access "
    "is refused \u2014 the default-deny principle.")

p_subsection(doc, "D. FIFO Message Deduplication")
p_body(doc,
    "Each FIFO queue maintains an internal map from DeduplicationID to the timestamp of last acceptance. If an identical ID arrives "
    "within five minutes, the queue returns the original MessageID and quietly drops the duplicate. Once the window passes, the same "
    "ID is accepted as fresh. Testing the exact boundary of this window \u2014 four minutes fifty-nine seconds versus five minutes "
    "one second \u2014 is trivial with FakeClock.")

p_subsection(doc, "E. Numeric-Aware Database Comparisons")
p_body(doc,
    "Scan filters in the database mocks run through a compareValues() function that first attempts to parse both operands as floats. "
    "If that works, comparison is numeric; if not, it falls back to lexicographic ordering. Without this, the string \"9\" would sort "
    "after \"10\" because the character '9' comes after '1'. A small detail, but one that has tripped up more than a few real-world "
    "query bugs.")

p_subsection(doc, "F. Dead-Letter Queues")
p_body(doc,
    "When creating a queue you can attach a dead-letter queue through DeadLetterConfig, specifying a TargetQueueURL and a "
    "MaxReceiveCount. After a message has been received more than MaxReceiveCount times without being deleted, the next call to "
    "ReceiveMessages quietly redirects it to the DLQ. This is the exact mechanism SQS, Service Bus, and Pub/Sub use to quarantine "
    "poison messages that consumers keep failing to process.")

p_subsection(doc, "G. Cost Simulation")
p_body(doc,
    "A cost.Tracker accumulates an estimated dollar amount for every operation executed. Default per-operation rates ship for all "
    "ten service categories, and SetRate() lets teams override them with their own pricing. Four accessors \u2014 TotalCost(), "
    "CostByService(), CostByOperation(), Reset() \u2014 give tests enough granularity to fail a CI build when a scenario would exceed "
    "a budget threshold. It is a rough estimate, not a billing replacement, but it catches the kind of runaway-loop mistakes that "
    "regular mocks silently ignore.")

p_subsection(doc, "H. Serverless Triggers")
p_body(doc,
    "Calling SetTrigger() wires a queue directly to a serverless function. From that point on, every message dropped onto the queue "
    "also gets dispatched to the function with the message body as input. This mirrors the SQS-to-Lambda, Service Bus-to-Azure "
    "Functions, and Pub/Sub-to-Cloud Functions pipelines that production systems rely on. RemoveTrigger() severs the connection.")

# V. IMPLEMENTATION
p_section(doc, "VI. IMPLEMENTATION")

p_subsection(doc, "A. Technology Choices")
p_body(doc,
    "Everything lives in one Go module (github.com/stackshy/cloudemu), built with Go 1.25. We import nothing outside the standard "
    "library\u2014not by accident but by deliberate constraint. Why Go specifically? Goroutines made concurrent tests natural to write. "
    "sync.RWMutex gave us fine-grained locking without pulling in a third-party lock library. Generics, which arrived in Go 1.18, "
    "meant the store layer could work with typed structs instead of interface{} everywhere. And frankly, Go compiles so fast that "
    "the full build barely registers as a line item in our CI pipeline.")

# >>> DIAGRAM: Implementation
add_img(doc, img_impl)

p_subsection(doc, "B. Generic In-Memory Store")
p_body(doc,
    "At the storage level, every mock delegates to memstore.Store[V]. It is, to be blunt, just a map[string]V behind a "
    "sync.RWMutex with four methods: Get, Set, Delete, All. We considered adding more\u2014range scans, conditional writes\u2014but "
    "decided against it. The simplicity pays off: because V is a generic type parameter, each service manipulates its own "
    "concrete struct without a single interface{} cast in sight.")

p_subsection(doc, "C. State Machine for Compute Lifecycle")
p_body(doc,
    "All three providers share a single state machine for VM lifecycle transitions. The valid paths look like this: "
    "pending \u2192 running \u2192 stopping \u2192 stopped \u2192 pending again for restarts; running \u2192 shutting-down \u2192 terminated, "
    "which is a one-way door; and running \u2192 restarting \u2192 running for reboots. Try to stop a VM that is already stopped and "
    "you get a FailedPrecondition error back. We debated whether to silently ignore invalid transitions the way some real providers "
    "do, but decided that failing loudly catches more bugs in user code.")

# >>> DIAGRAM: VM lifecycle
add_img(doc, img_vm)

p_subsection(doc, "D. ID Generation")
p_body(doc, "AWS: arn:aws:<service>:::<resource>")
p_body(doc, "Azure: /subscriptions/<sub>/resourceGroups/<rg>/providers/<type>/<resource>")
p_body(doc, "GCP: projects/<project>/<collection>/<resource>")

p_subsection(doc, "E. Cross-Cutting Concerns")
p_body(doc,
    "The portable API layer wraps four optional components around every driver call. First, a Recorder that timestamps each API "
    "invocation\u2014tests can later assert that calls happened in a particular order. Second, a Metrics Collector counting invocations "
    "and measuring latency per operation; we added it after noticing some of our own tests were making more calls than we realized. "
    "Third, a Rate Limiter built on a token bucket. Exhaust the tokens and the next call returns Throttled, which is exactly the "
    "error your retry logic needs to handle. Fourth, an Error Injector: tell it to fail on the third call to GetItem and it will, "
    "no questions asked. We use this one a lot when testing error-recovery paths.")

p_subsection(doc, "F. Code Size")
p_body(doc,
    "At the time of writing the codebase contains roughly 13,850 lines of Go spread over 78 files, tests excluded. The test code adds another 1,600 lines.")

# VI. EVALUATION
p_section(doc, "VII. EVALUATION")

p_subsection(doc, "A. Correctness")
p_body(doc, "The test suite contains 31 tests. We group them into three tiers, roughly in order of how much infrastructure each one exercises:")
p_body(doc,
    "Tier 1 \u2014 Service-Level (8 tests). Storage lifecycle with pagination, compute state-machine transitions including illegal "
    "ones, database CRUD plus query and scan operators, and a cross-provider portability test that replays the same scenario on "
    "AWS, Azure, and GCP side by side.",
    indent_first=True)
p_body(doc,
    "Tier 2 \u2014 Cross-Cutting (4 tests). One test per wrapper: forced error injection, token-bucket rate limiting, API call "
    "recording, and per-operation metric collection.",
    indent_first=True)
p_body(doc, "Tier 3 \u2014 Realistic Behaviors (19 tests), detailed in Table IV.", indent_first=True)

table_caption(doc, "TABLE IV. REALISTIC BEHAVIOR TEST SUITE")
add_ieee_table(doc,
    ["Test", "Behavior", "Prov."],
    [
        ["AutoMetricGen", "VM emits 5\u00d75 metrics", "All 3"],
        ["AlarmAutoEval", "Threshold fires alarm", "All 3"],
        ["AlarmByMetrics", "VM\u2192metrics\u2192alarm", "All 3"],
        ["StopEmitsZero", "Stop zeros, alarm", "All 3"],
        ["StartEmitsRun", "Start running vals", "All 3"],
        ["IAMPermission", "Policy+wildcards", "All 3"],
        ["FIFODedup", "5-min dedup", "All 3"],
        ["NumericCmp", "Numeric DB cmp", "All 3"],
        ["ScanOps", "\u2264,\u2265,BEGINS_WITH", "All 3"],
        ["DeadLetterQ", "Poison routing", "All 3"],
        ["CostTracker", "Op cost tally", "All"],
        ["ServerlessTrig", "Queue\u2192function", "All 3"],
    ],
    col_widths=[1500, 1900, 800])

p_body(doc,
    "All 31 tests pass on every run. Not once have we hit a flaky failure, which is a sharp contrast to the integration suites we "
    "maintained before building CloudEmu. We think two factors explain the consistency: state never leaves the process (no disk, no "
    "network), and time is fake\u2014so there is nothing for race conditions to latch onto.")

p_subsection(doc, "B. Integration Testing")
p_body(doc,
    "Beyond the unit tests, we wrote 6 HTTP-level integration tests inside a companion project called ZipZopCloud that exposes "
    "CloudEmu through REST endpoints. Those tests send real HTTP requests, parse real JSON responses, and confirm that the behavioral "
    "guarantees survive the extra layer of serialization and routing.")

p_subsection(doc, "C. Performance")
p_body(doc,
    "Individual operations complete in microseconds because they amount to map lookups guarded by a mutex. There is no container "
    "to start, no port to bind, no TCP handshake to complete. On a mid-range laptop the entire 31-test suite finishes in well "
    "under half a second.")

p_subsection(doc, "D. Comparison with Existing Tools")

table_caption(doc, "TABLE V. COMPARISON WITH EXISTING TOOLS")
add_ieee_table(doc,
    ["", "Ours", "LStack", "Moto", "Azurite", "FBEm"],
    [
        ["Lang", "Go", "Py", "Py", "Node", "Node"],
        ["Cloud", "3", "1", "1", "1", "1"],
        ["Docker", "No", "Yes", "No", "No", "No"],
        ["Start", "0ms", "3s", "500ms", "2s", "4s"],
        ["InProc", "Yes", "No", "Yes", "No", "No"],
        ["Fidel.", "High", "High", "Med", "Med", "Med"],
        ["Svcs", "30", "70+", "~50", "3", "5"],
        ["Deps", "0", "2", "1", "1", "2"],
    ],
    col_widths=[650, 650, 700, 700, 750, 700])

# >>> DIAGRAM: Comparison bar chart
add_img(doc, img_bar)

# >>> DIAGRAM: Radar
add_img(doc, img_radar)

# VII. DISCUSSION AND LIMITATIONS
p_section(doc, "VIII. DISCUSSION AND LIMITATIONS")

p_subsection(doc, "A. Limitations")
p_body(doc,
    "Several services are shallow. The S3 mock, to take one example, lacks object versioning, lifecycle transitions, and "
    "cross-region replication. We deliberately traded depth for breadth \u2014 ten categories across three providers \u2014 in this "
    "initial release.",
    indent_first=True)
p_body(doc,
    "There is no network simulation whatsoever. Partition failures, DNS lookup errors, TLS negotiation issues \u2014 none of those "
    "can be reproduced because every call is a direct Go function invocation inside the same process.",
    indent_first=True)
p_body(doc,
    "Right now the library is accessible only from Go. Teams writing Python, Java, or TypeScript would need the HTTP wrapper "
    "described in Future Work before they could use it.",
    indent_first=True)
p_body(doc,
    "IAM emulation covers action and resource matching with wildcard patterns, but condition keys, permission boundaries, and "
    "service control policies remain unimplemented. Real IAM is enormously complex; we covered the twenty percent of it that "
    "handles eighty percent of test scenarios.",
    indent_first=True)

p_subsection(doc, "B. Threats to Validity")
p_body(doc,
    "Every correctness claim we make rests on how faithfully we interpreted the official documentation. We never ran conformance "
    "checks against live AWS, Azure, or GCP endpoints, so subtle discrepancies are almost certainly lurking \u2014 especially around "
    "undocumented edge cases or region-specific behaviors. The 31 tests validate what we intentionally built, not the full surface "
    "area of any real cloud service.")

p_subsection(doc, "C. Future Work")
p_bullet(doc, "Deepen individual services. S3 versioning, Lambda cold-start behavior, and richer database query operators are at the top of our list.")
p_bullet(doc, "Stand up an HTTP facade so that Python, Java, and TypeScript projects can talk to CloudEmu the same way they talk to LocalStack\u2019s AWS-compatible endpoints.")
p_bullet(doc, "Create a conformance harness that runs identical assertions against both CloudEmu and real cloud APIs, flagging behavioral drift before it surprises users.")
p_bullet(doc, "Publish rigorous benchmarks comparing startup and per-operation latency against Docker-based emulators, particularly in CI/CD pipelines where container spin-up dominates wall time.")

# VIII. CONCLUSION AND FUTURE WORK
p_section(doc, "IX. CONCLUSION")
p_body(doc,
    "CloudEmu started because we were frustrated with the state of multi-cloud testing in Go. What came out of that frustration is "
    "a library that provides in-memory fakes for 10 services across AWS, Azure, and GCP, structured in three layers that keep "
    "portable APIs, driver contracts, and provider backends from tangling together. Eight behavioral emulations raise the bar "
    "above simple record storage: compute lifecycle metrics, threshold-based alarm evaluation, IAM policy walks with wildcards, "
    "FIFO deduplication, numeric-aware database comparisons, dead-letter queue routing, per-operation cost estimation, and "
    "serverless event-source dispatch.")
p_body(doc,
    "To our knowledge, nothing else on the market bundles multi-cloud support, in-process execution, zero external dependencies, "
    "and this level of behavioral fidelity into one package. The 31-test suite runs in under 500 ms, starts instantly, and has "
    "not produced a flaky result in hundreds of CI runs. The code is open-source at github.com/stackshy/cloudemu. If it saves "
    "another team even half the debugging hours it saved us, it will have been worth the effort.",
    indent_first=True)

# REFERENCES — keep_with_next on heading so it stays with first ref
p = p_section(doc, "REFERENCES")
p.paragraph_format.keep_with_next = True
refs = [
    '[1]  W. Hummer et al., "LocalStack: A fully functional local cloud stack," GitHub, 2024.',
    '[2]  S. Garnaat et al., "Moto: Mock AWS Services," GitHub, 2024.',
    '[3]  P. Garousi et al., "A systematic review on cloud testing," ACM Comput. Surv., vol. 52, no. 4, 2019.',
    '[4]  R. S. Kumar, "Cloud-native performance testing in CI/CD pipelines," ResearchGate, 2024.',
    '[5]  G. Meszaros, xUnit Test Patterns. Addison-Wesley, 2007.',
    '[6]  M. Fowler, "TestDouble," martinfowler.com, 2006.',
    '[7]  Microsoft, "Azurite open-source emulator," MS Learn, 2024.',
    '[8]  Google, "Firebase Local Emulator Suite," Firebase docs, 2024.',
    '[9]  Google, "Go Cloud Development Kit," GitHub, 2024.',
    '[10] Apache, "Apache Libcloud," 2024.',
    '[11] N. Goonasekera et al., "CloudBridge," Proc. XSEDE16, ACM, 2016.',
    '[12] R. N. Calheiros et al., "CloudSim," Softw. Pract. Exper., vol. 41, no. 1, pp. 23\u201350, 2011.',
    '[13] M. C. Silva Filho et al., "CloudSim Plus," Proc. IFIP/IEEE IM, pp. 400\u2013406, 2017.',
    '[14] R. Buyya et al., "CloudSim 7G," Softw. Pract. Exper., 2025.',
    '[15] Apache, "Apache jclouds," 2024.',
    '[16] T. Roome, "WireMock," wiremock.io, 2024.',
    '[17] H. Wang et al., "Mocking frameworks in Apache," Empir. Softw. Eng., vol. 28, no. 6, 2023.',
]
for i, ref in enumerate(refs):
    is_last = (i == len(refs) - 1)
    p_ref(doc, ref, keep_with_next=not is_last)

# APPENDIX
p_section(doc, "APPENDIX: CROSS-PROVIDER METRIC MAPPING")

table_caption(doc, "TABLE VI. CROSS-PROVIDER METRIC MAPPING")
add_ieee_table(doc,
    ["Metric", "AWS", "Azure", "GCP"],
    [
        ["CPU", "CPUUtil=25", "%CPU=25", "cpu/util=0.25"],
        ["NetIn", "NetIn=1024", "NetInTot=1024", "recv_bytes=1024"],
        ["NetOut", "NetOut=512", "NetOutTot=512", "sent_bytes=512"],
        ["DiskRd", "DiskRdOps=100", "DiskRd/s=100", "rd_ops=100"],
        ["DiskWr", "DiskWrOps=50", "DiskWr/s=50", "wr_ops=50"],
    ],
    col_widths=[750, 1250, 1350, 1250])

doc.save(OUTPUT)
print(f"\nDone! Saved to {OUTPUT}")
print("Format: IEEE two-column, Times New Roman, black text, italic sub-headings")
print("9 diagrams inserted at appropriate locations.")
